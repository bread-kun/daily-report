# @athor	Asan
# see more: https://python-docx.readthedocs.io/en/latest/index.html
from docx import Document
from docx.shared import Inches
import time
import re
from git import Repo as rp
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
	def __cell_format__(tc):
		tc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
		tc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		_tb.rows[0].height = Inches(0.8)

	def __each__(array, fn):
		for n in array:
			fn(n)
	def __set_rows__(row, contents):
		for i,c in enumerate(contents):
			row.cells[i].text = contents[i]

	doc = Document()
	doc.add_paragraph('(个人)');
	_tb = doc.add_table(rows=1, cols=4, style='Table Grid')
	_tb.alignment = WD_TABLE_ALIGNMENT.CENTER
	_first_cell = _tb.rows[0].cells[0]
	_first_cell.merge(_tb.rows[0].cells[3])
	_first_cell.text = "工作日报"
	_first_cell.paragraphs[-1].style = 'Heading 1'
	__cell_format__(_tb.rows[0].cells[0])
	_tb.rows[0].height = Inches(0.8)


	_row = _tb.add_row()
	__set_rows__(_row, ['姓名','xxx','日期',time.strftime("%Y-%m-%d", time.localtime())])
	__each__(_row.cells, __cell_format__)
	_row.height = Inches(0.5)

	# 多库支持
	_muti_repo = [
				rp('D:/repo1'),
				rp('D:/repo2'),
			]
	_history = []
	# 时间为一天内
	for _rp in _muti_repo:
		_history.extend(_rp.git.log(r'--pretty=format:"%ad %an -> %cd #coms: %s\\ufe48"',r'--author=yourGitAuthorName', r'--date=format:"%Y-%m-%d %H:%M"', r'--since=1.day').split(r'\\ufe48'))

	_row = _tb.add_row()
	_row.cells[0].text = '今日工作汇报'
	_row.cells[1].merge(_row.cells[3])
	__cell_format__(_row.cells[0])
	_t_count = 0
	for i, cmt in enumerate(_history):
		if len(cmt.strip()) > 10:
			_content = cmt.split(r'#coms: ')[1]
			# 过滤纯数字英文的提交内容，至少要有说明，不然默认是无信息提交
			if re.match('^[a-zA-Z0-9 \'\"\:/_\.]*$', _content):
				continue
			for _each in _content.split(chr(32)):
				_t_count += 1
				if _t_count == 1:
					_row.cells[1].text = '{0}.{1}'.format(_t_count,_each)
				else:
					_row.cells[1].add_paragraph('{0}.{1}'.format(_t_count,_each))
	_row.height = Inches(0.5)

	_row = _tb.add_row()
	_row.cells[1].merge(_row.cells[3])
	__set_rows__(_row, ['明日工作计划', '设定明日工作计划'])
	_row.cells[1].paragraphs[-1].style = 'List Number'
	_row.height = Inches(0.5)
	__cell_format__(_row.cells[0])

	last_row = _tb.add_row()
	last_row.cells[0].merge(last_row.cells[3])
	last_row.cells[0].text = "★工作改进、优化、建议、意见或今天的收获、心得"
	last_row.height = Inches(0.8)

	_tb.columns[0].width = Inches(0.5)

	doc.save('xxx工作日报表-%s.docx' %time.strftime("%Y-%m-%d", time.localtime()))

main()